import os, uuid, io
from flask import Flask, request, render_template, send_file, jsonify
from werkzeug.utils import secure_filename
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
import PyPDF2
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import qr

app = Flask(__name__)
UPLOAD, OUTPUT = "uploads", "outputs"
ALLOWED = {"pdf", "docx"}

os.makedirs(UPLOAD, exist_ok=True)
os.makedirs(OUTPUT, exist_ok=True)


def allowed(name):
    return "." in name and name.rsplit(".", 1)[1].lower() in ALLOWED


def unique(folder, name):
    base, ext = os.path.splitext(name)
    return os.path.join(folder, f"{base}_{uuid.uuid4().hex[:8]}{ext}")


def add_field(run, field):
    for tag, val in [("w:fldChar", "begin"), ("w:instrText", field), ("w:fldChar", "end")]:
        el = OxmlElement(tag)
        if tag == "w:fldChar":
            el.set(qn("w:fldCharType"), val)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = val
        run._r.append(el)


def style(run):
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_pdf_footer(src, dst, text):
    reader, writer = PyPDF2.PdfReader(src), PyPDF2.PdfWriter()
    total = len(reader.pages)

    for i, page in enumerate(reader.pages, 1):
        w, h = float(page.mediabox.width), float(page.mediabox.height)
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=(w, h))

        c.setStrokeColor(HexColor("#AAAAAA"))
        c.setLineWidth(0.5)
        c.line(50, 38, w - 50, 38)

        c.setFont("Helvetica", 9)
        c.setFillColor(HexColor("#555555"))
        c.drawCentredString(w / 2, 22, f"{text} | Page {i} of {total}")
        c.save()

        packet.seek(0)
        page.merge_page(PyPDF2.PdfReader(packet).pages[0])
        writer.add_page(page)

    with open(dst, "wb") as f:
        writer.write(f)


def add_docx_footer(src, dst, text):
    doc = Document(src)

    for section in doc.sections:
        section.different_first_page_header_footer = False
        para = section.footer.paragraphs[0]
        para.clear()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for item in [f"{text} | Page ", "PAGE", " of ", "NUMPAGES"]:
            run = para.add_run("" if item in {"PAGE", "NUMPAGES"} else item)
            style(run)
            if item in {"PAGE", "NUMPAGES"}:
                add_field(run, item)

    doc.save(dst)


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/process", methods=["POST"])
def process():
    file = request.files.get("file")
    footer = request.form.get("footer_text", "").strip()

    if not file or file.filename == "":
        return jsonify(error="No file selected."), 400
    if not footer:
        return jsonify(error="Footer text cannot be empty."), 400
    if not allowed(file.filename):
        return jsonify(error="Only PDF and DOCX files are supported."), 400

    name = secure_filename(file.filename)
    ext = name.rsplit(".", 1)[1].lower()

    src = unique(UPLOAD, name)
    out_name = name.replace(f".{ext}", f"_with_footer.{ext}")
    dst = unique(OUTPUT, out_name)

    file.save(src)

    try:
        {"pdf": add_pdf_footer, "docx": add_docx_footer}[ext](src, dst, footer)
    except Exception as e:
        return jsonify(error=f"Processing failed: {e}"), 500

    mime = "application/pdf" if ext == "pdf" else \
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return send_file(dst, as_attachment=True, download_name=out_name, mimetype=mime)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 5000)))
